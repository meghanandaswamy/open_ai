import os
import glob
import docx
import pandas as pd
from dotenv import load_dotenv
from typing import List
import openai
import tiktoken

from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

# Load API key
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# Paths
DOCS_FOLDER = "docs"
MENU_EXCEL_PATH = r"C:\Users\Meghananda\PycharmProjects\PythonProject4\docs\finacle_menus.xlsx"
CHROMA_PATH = "rag_store"

# Tokenizer
tokenizer = tiktoken.get_encoding("cl100k_base")  # GPT-4 tokenizer


def load_word_documents(folder_path: str) -> List[Document]:
    documents = []
    for file in glob.glob(os.path.join(folder_path, "*.docx")):
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        if text:
            documents.append(Document(page_content=text, metadata={"source": file}))
    return documents


def load_menu_names(excel_path: str) -> List[str]:
    df = pd.read_excel(excel_path)
    return df.iloc[:, 0].dropna().astype(str).tolist()


def chunk_documents(docs: List[Document], chunk_size=800, chunk_overlap=100):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(docs)


def build_vectorstore(chunks: List[Document]):
    if os.path.exists(CHROMA_PATH):
        db = Chroma(persist_directory=CHROMA_PATH, embedding_function=OpenAIEmbeddings())
    else:
        db = Chroma.from_documents(chunks, OpenAIEmbeddings(), persist_directory=CHROMA_PATH)
        db.persist()
    return db


def get_token_limited_context(chunks: List[str], max_tokens: int) -> str:
    total_tokens = 0
    final_chunks = []
    for chunk in chunks:
        tokens = len(tokenizer.encode(chunk))
        if total_tokens + tokens <= max_tokens:
            final_chunks.append(chunk)
            total_tokens += tokens
        else:
            break
    return "\n".join(final_chunks)


def generate_test_scenarios(menu: str, db, max_context_tokens=800) -> str:
    search = db.similarity_search(menu, k=5)
    raw_chunks = [doc.page_content for doc in search]
    context = get_token_limited_context(raw_chunks, max_context_tokens)

    prompt = f"""
You are a senior test engineer for the Finacle core banking application.

Based on the following internal documentation:
\"\"\"
{context}
\"\"\"

Generate **3 real-world test scenarios** for the Finacle menu named **"{menu}"**:
1. A positive test scenario covering the main business path.
2. Another positive test covering an alternate valid path.
3. A negative test scenario handling invalid or edge input.

Give the output in a clear bullet-point list.
"""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",  # or "gpt-3.5-turbo" for cost saving
            messages=[
                {"role": "system", "content": "You are an expert in banking software QA."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=600
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"[ERROR]: {e}"


if __name__ == "__main__":
    print("🔍 Loading documents...")
    docs = load_word_documents(DOCS_FOLDER)
    chunks = chunk_documents(docs)
    db = build_vectorstore(chunks)

    menus = load_menu_names(MENU_EXCEL_PATH)

    print("\n📋 Available Menus:")
    for idx, name in enumerate(menus, 1):
        break
        # print(f"{idx}. {name}")

    print("\nEnter menu name (or type 'exit'):")
    while True:
        menu = input("🔸 Menu: ").strip()
        if menu.lower() == "exit":
            break
        if menu not in menus:
            print("❌ Menu not found. Please enter a valid menu from the list.")
            continue

        print(f"\n📄 Generating test scenarios for: {menu}")
        result = generate_test_scenarios(menu, db)
        print(result)
        print("-" * 60)
